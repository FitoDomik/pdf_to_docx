import os
import sys
import tempfile
import re
import uuid
from PIL import Image
from PIL.ImageQt import ImageQt
import numpy as np
import torch
import fitz  
from pdf2image import convert_from_path
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QPushButton, 
                           QFileDialog, QProgressBar, QLabel, QHBoxLayout,
                           QMessageBox, QComboBox, QGroupBox, QGridLayout)
from PyQt6.QtCore import Qt, QMimeData, QUrl, pyqtSignal, QThread
from PyQt6.QtGui import QDragEnterEvent, QDropEvent, QPixmap, QIcon
try:
    from paddleocr import PaddleOCR
    PADDLE_OCR_AVAILABLE = True
except ImportError:
    print("PaddleOCR не установлен. Будет доступен только TrOCR.")
    PADDLE_OCR_AVAILABLE = False
try:
    from transformers import TrOCRProcessor, VisionEncoderDecoderModel
    TROCR_AVAILABLE = True
except ImportError:
    print("TrOCR не установлен. Будет доступен только PaddleOCR.")
    TROCR_AVAILABLE = False
if not PADDLE_OCR_AVAILABLE and not TROCR_AVAILABLE:
    print("ОШИБКА: Ни один из OCR движков (PaddleOCR, TrOCR) не установлен.")
    print("Установите хотя бы один из них:")
    print("pip install paddleocr paddlepaddle")
    print("pip install transformers")
    sys.exit(1)
def get_file_extension(file_path):
    _, ext = os.path.splitext(file_path)
    return ext[1:] if ext else ""
def create_thumbnail(image_path, max_size=400):
    image = Image.open(image_path)
    width, height = image.size
    if width > height:
        new_width = max_size
        new_height = int(height * max_size / width)
    else:
        new_height = max_size
        new_width = int(width * max_size / height)
    thumbnail = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
    return thumbnail
def is_likely_header(text):
    if len(text) > 100:
        return False
    if text.strip().endswith(':'):
        return True
    if text.isupper():
        return True
    if re.match(r'^\d+\.?\s', text):
        return True
    return False
def is_list_item(text):
    if text.strip().startswith(('•', '-', '*', '○', '▪', '▫')):
        return True
    if re.match(r'^\d+\.?\s', text):
        return True
    return False
def clean_temp_files():
    temp_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "temp")
    if os.path.exists(temp_dir):
        for file_name in os.listdir(temp_dir):
            file_path = os.path.join(temp_dir, file_name)
            try:
                if os.path.isfile(file_path):
                    os.unlink(file_path)
            except Exception as e:
                print(f"Ошибка при удалении {file_path}: {e}")
def detect_language(text):
    if re.search(r'[а-яА-ЯёЁ]', text):
        return 'ru'
    return 'en'
def convert_pdf_to_images(pdf_path, dpi=300, pages=None):
    temp_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "temp")
    os.makedirs(temp_dir, exist_ok=True)
    file_prefix = f"pdf_page_{uuid.uuid4().hex}_"
    images = convert_from_path(
        pdf_path, 
        dpi=dpi, 
        output_folder=temp_dir,
        fmt="png",
        output_file=file_prefix,
        paths_only=True,
        first_page=pages[0] + 1 if pages else None,
        last_page=pages[-1] + 1 if pages else None
    )
    return images
def extract_images_from_pdf(pdf_path):
    pdf_document = fitz.open(pdf_path)
    images_data = []
    temp_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "temp")
    os.makedirs(temp_dir, exist_ok=True)
    for page_idx, page in enumerate(pdf_document):
        image_list = page.get_images(full=True)
        for img_idx, img_info in enumerate(image_list):
            xref = img_info[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            image_path = os.path.join(temp_dir, f"pdf_img_{page_idx}_{img_idx}_{os.getpid()}.{image_ext}")
            with open(image_path, 'wb') as img_file:
                img_file.write(image_bytes)
            rect = page.get_image_bbox(img_info)
            images_data.append({
                'page': page_idx + 1,
                'path': image_path,
                'rect': [rect.x0, rect.y0, rect.x1, rect.y1]
            })
    return images_data 
def save_docx_file(recognized_data, output_path):
    doc = docx.Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    for page_data in recognized_data:
        text_blocks = sorted(page_data['text_blocks'], 
                            key=lambda block: block['coordinates'][0][1])
        for block in text_blocks:
            text = block['text']
            block_type = block.get('type', 'text')
            if block_type == 'header':
                p = doc.add_paragraph()
                p.style = 'Heading 1'
                p.add_run(text)
            else:
                p = doc.add_paragraph()
                p.style = 'Normal'
                p.add_run(text)
                if text.strip().startswith(('•', '-', '*')):
                    p.style = 'List Bullet'
                elif text.strip()[0].isdigit() and text.strip()[1:].startswith('. '):
                    p.style = 'List Number'
        if page_data != recognized_data[-1]:
            doc.add_page_break()
    doc.save(output_path)
    return output_path
def extract_text_from_pdf(pdf_path):
    pdf_document = fitz.open(pdf_path)
    text_blocks = []
    for page_idx, page in enumerate(pdf_document):
        blocks = page.get_text("blocks")
        for block in blocks:
            text_blocks.append({
                'page': page_idx + 1,
                'text': block[4],
                'coordinates': [[block[0], block[1]], [block[2], block[1]], 
                               [block[2], block[3]], [block[0], block[3]]]
            })
    return text_blocks
class OCRProcessor:
    def __init__(self):
        self.models_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "models")
        os.makedirs(self.models_dir, exist_ok=True)
        self.paddle_ocr_initialized = False
        self.trocr_initialized = False
        self.paddle_ocr = None
        self.trocr_processor = None
        self.trocr_model = None
    def init_paddle_ocr(self):
        if not PADDLE_OCR_AVAILABLE:
            raise Exception("PaddleOCR не установлен. Установите его с помощью: pip install paddleocr paddlepaddle")
        if not self.paddle_ocr_initialized:
            det_model_dir = os.path.join(self.models_dir, "paddle_det")
            rec_model_dir = os.path.join(self.models_dir, "paddle_rec")
            cls_model_dir = os.path.join(self.models_dir, "paddle_cls")
            os.makedirs(det_model_dir, exist_ok=True)
            os.makedirs(rec_model_dir, exist_ok=True)
            os.makedirs(cls_model_dir, exist_ok=True)
            try:
                print("Попытка инициализации PaddleOCR с указанными путями...")
                self.paddle_ocr = PaddleOCR(
                    use_angle_cls=True,
                    lang='ru',
                    det_model_dir=det_model_dir,
                    rec_model_dir=rec_model_dir,
                    cls_model_dir=cls_model_dir
                )
                self.paddle_ocr_initialized = True
                print("PaddleOCR успешно инициализирован с указанными путями")
            except Exception as e:
                print(f"Ошибка при инициализации PaddleOCR с указанными путями: {str(e)}")
                print("Пробуем инициализировать PaddleOCR без указания путей...")
                try:
                    self.paddle_ocr = PaddleOCR(
                        use_angle_cls=True,
                        lang='ru'
                    )
                    self.paddle_ocr_initialized = True
                    print("PaddleOCR успешно инициализирован без указания путей")
                except Exception as e:
                    error_msg = f"Ошибка при инициализации PaddleOCR: {str(e)}"
                    print(error_msg)
                    raise Exception(error_msg)
    def init_trocr(self):
        if not TROCR_AVAILABLE:
            raise Exception("TrOCR не установлен. Установите его с помощью: pip install transformers")
        if not self.trocr_initialized:
            cache_dir = os.path.join(self.models_dir, "trocr")
            os.makedirs(cache_dir, exist_ok=True)
            os.environ['TRANSFORMERS_CACHE'] = cache_dir
            try:
                print("Загрузка модели TrOCR...")
                self.trocr_processor = TrOCRProcessor.from_pretrained(
                    'microsoft/trocr-base-handwritten',
                    cache_dir=cache_dir
                )
                self.trocr_model = VisionEncoderDecoderModel.from_pretrained(
                    'microsoft/trocr-base-handwritten',
                    cache_dir=cache_dir
                )
                if torch.cuda.is_available():
                    self.trocr_model = self.trocr_model.to('cuda')
                self.trocr_initialized = True
                print("Модель TrOCR успешно загружена")
            except Exception as e:
                error_msg = f"Ошибка инициализации TrOCR: {str(e)}"
                print(error_msg)
                raise Exception(error_msg) 
    def process_images(self, image_paths, progress_signal=None, ocr_engine="PaddleOCR"):
        recognized_data = []
        try:
            if ocr_engine == "PaddleOCR":
                self.init_paddle_ocr()
                process_func = self._process_with_paddleocr
            else:  
                self.init_trocr()
                process_func = self._process_with_trocr
            total_images = len(image_paths)
            for i, image_path in enumerate(image_paths):
                try:
                    page_data = process_func(image_path)
                    recognized_data.append(page_data)
                    if progress_signal:
                        progress = 10 + int(70 * (i + 1) / total_images)
                        progress_signal.emit(progress)
                except Exception as e:
                    error_msg = f"Ошибка при обработке изображения {image_path}: {str(e)}"
                    print(error_msg)
                    recognized_data.append({
                        'image_path': image_path,
                        'text_blocks': [{
                            'text': f"ОШИБКА РАСПОЗНАВАНИЯ: {str(e)}",
                            'confidence': 0.0,
                            'coordinates': [[0, 0], [100, 0], [100, 100], [0, 100]],
                            'type': 'text'
                        }]
                    })
                    if progress_signal:
                        progress = 10 + int(70 * (i + 1) / total_images)
                        progress_signal.emit(progress)
        except Exception as e:
            error_msg = f"Ошибка при инициализации OCR движка {ocr_engine}: {str(e)}"
            print(error_msg)
            raise Exception(error_msg)
        return recognized_data
    def _process_with_paddleocr(self, image_path):
        result = self.paddle_ocr.ocr(image_path)
        page_data = {
            'image_path': image_path,
            'text_blocks': []
        }
        try:
            if isinstance(result, list) and len(result) > 0:
                ocr_result = result[0]
                if hasattr(ocr_result, 'rec_texts') and hasattr(ocr_result, 'rec_polys'):
                    texts = ocr_result.rec_texts
                    polys = ocr_result.rec_polys
                    scores = ocr_result.rec_scores if hasattr(ocr_result, 'rec_scores') else [0.9] * len(texts)
                    for i, (text, poly, score) in enumerate(zip(texts, polys, scores)):
                        text_block = {
                            'text': text,
                            'confidence': float(score),
                            'coordinates': poly.tolist() if hasattr(poly, 'tolist') else poly,
                            'type': 'header' if len(text) < 50 and float(score) > 0.9 else 'text'
                        }
                        page_data['text_blocks'].append(text_block)
                elif isinstance(ocr_result, dict):
                    if 'rec_texts' in ocr_result and 'rec_polys' in ocr_result:
                        texts = ocr_result['rec_texts']
                        polys = ocr_result['rec_polys']
                        scores = ocr_result.get('rec_scores', [0.9] * len(texts))
                        for i, (text, poly, score) in enumerate(zip(texts, polys, scores)):
                            text_block = {
                                'text': text,
                                'confidence': float(score),
                                'coordinates': poly.tolist() if hasattr(poly, 'tolist') else poly,
                                'type': 'header' if len(text) < 50 and float(score) > 0.9 else 'text'
                            }
                            page_data['text_blocks'].append(text_block)
                else:
                    for line in ocr_result:
                        if isinstance(line, tuple) and len(line) == 2:
                            coords = line[0]
                            text, confidence = line[1]
                        elif isinstance(line, list) and len(line) >= 2:
                            coords = line[0]
                            text = line[1]
                            confidence = line[2] if len(line) > 2 else 0.9
                        elif isinstance(line, dict):
                            coords = line.get('box', [[0, 0], [100, 0], [100, 100], [0, 100]])
                            text = line.get('text', '')
                            confidence = line.get('confidence', 0.9)
                        else:
                            print(f"Неизвестный формат результата: {line}")
                            continue
                        text_block = {
                            'text': text,
                            'confidence': float(confidence) if isinstance(confidence, (int, float)) else 0.9,
                            'coordinates': coords,
                            'type': 'header' if len(text) < 50 and float(confidence) > 0.9 else 'text'
                        }
                        page_data['text_blocks'].append(text_block)
        except Exception as e:
            print(f"Ошибка при обработке результата PaddleOCR: {str(e)}")
            page_data['text_blocks'].append({
                'text': f"ОШИБКА РАСПОЗНАВАНИЯ: {str(e)}",
                'confidence': 0.0,
                'coordinates': [[0, 0], [100, 0], [100, 100], [0, 100]],
                'type': 'text'
            })
        return page_data
    def _process_with_trocr(self, image_path):
        image = Image.open(image_path).convert("RGB")
        pixel_values = self.trocr_processor(image, return_tensors="pt").pixel_values
        if torch.cuda.is_available():
            pixel_values = pixel_values.to('cuda')
        generated_ids = self.trocr_model.generate(pixel_values)
        generated_text = self.trocr_processor.batch_decode(generated_ids, skip_special_tokens=True)[0]
        page_data = {
            'image_path': image_path,
            'text_blocks': [{
                'text': generated_text,
                'confidence': 0.9,  
                'coordinates': [[0, 0], [image.width, 0], [image.width, image.height], [0, image.height]],
                'type': 'text'
            }]
        }
        return page_data
    def analyze_document_structure(self, recognized_data):
        document_structure = {
            'pages': []
        }
        for page_idx, page_data in enumerate(recognized_data):
            page_structure = {
                'page_number': page_idx + 1,
                'elements': []
            }
            text_blocks = sorted(page_data['text_blocks'], 
                                key=lambda block: block['coordinates'][0][1])
            for block in text_blocks:
                element_type = self._determine_element_type(block)
                element = {
                    'type': element_type,
                    'text': block['text'],
                    'coordinates': block['coordinates']
                }
                page_structure['elements'].append(element)
            document_structure['pages'].append(page_structure)
        return document_structure
    def _determine_element_type(self, text_block):
        text = text_block['text']
        if len(text) < 50 and text.strip().isupper():
            return 'heading1'
        elif len(text) < 100 and text.strip().endswith(':'):
            return 'heading2'
        if text.strip().startswith(('•', '-', '*', '1.', '2.', '3.')):
            return 'list_item'
        return 'paragraph' 
class ProcessingThread(QThread):
    progress_signal = pyqtSignal(int)
    finished_signal = pyqtSignal(str)
    error_signal = pyqtSignal(str)
    def __init__(self, ocr_processor, file_paths, output_path, ocr_engine):
        super().__init__()
        self.ocr_processor = ocr_processor
        self.file_paths = file_paths
        self.output_path = output_path
        self.ocr_engine = ocr_engine
    def run(self):
        try:
            all_images = []
            for file_path in self.file_paths:
                ext = get_file_extension(file_path)
                if ext.lower() == 'pdf':
                    pdf_images = convert_pdf_to_images(file_path)
                    all_images.extend(pdf_images)
                    self.progress_signal.emit(10)
                else:
                    all_images.append(file_path)
                    self.progress_signal.emit(10)
            recognized_data = self.ocr_processor.process_images(
                all_images, 
                self.progress_signal,
                self.ocr_engine
            )
            self.progress_signal.emit(80)
            docx_path = save_docx_file(recognized_data, self.output_path)
            self.progress_signal.emit(100)
            self.finished_signal.emit(docx_path)
        except Exception as e:
            self.error_signal.emit(f"Ошибка при обработке: {str(e)}")
class DropArea(QWidget):
    def __init__(self, parent):
        super().__init__()
        self.parent = parent
        self.setAcceptDrops(True)
        self.setMinimumHeight(200)
        self.setStyleSheet("border: 2px dashed #cccccc; border-radius: 5px;")
        layout = QVBoxLayout(self)
        icon_label = QLabel()
        icon = parent.style().standardIcon(parent.style().StandardPixmap.SP_FileIcon)
        icon_label.setPixmap(icon.pixmap(64, 64))
        icon_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        icon_label.setStyleSheet("border: none;")
        layout.addWidget(icon_label)
        self.label = QLabel("Перетащите файлы сюда")
        self.label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.label.setStyleSheet("border: none;")
        layout.addWidget(self.label)
    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
            self.setStyleSheet("background-color: #e0f7fa; border: 2px dashed #00acc1; border-radius: 5px;")
    def dragLeaveEvent(self, event):
        self.setStyleSheet("border: 2px dashed #cccccc; border-radius: 5px;")
    def dropEvent(self, event: QDropEvent):
        self.setStyleSheet("border: 2px dashed #cccccc; border-radius: 5px;")
        file_paths = []
        mime_data: QMimeData = event.mimeData()
        if mime_data.hasUrls():
            for url in mime_data.urls():
                file_path = url.toLocalFile()
                ext = get_file_extension(file_path)
                if ext.lower() in ['pdf', 'jpg', 'jpeg', 'png', 'bmp']:
                    file_paths.append(file_path)
        if file_paths:
            self.parent.file_paths = file_paths
            self.parent.update_convert_button()
            self.label.setText(f"Выбрано файлов: {len(file_paths)}")
        event.acceptProposedAction()
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.ocr_processor = OCRProcessor()
        self.file_paths = []
        self.output_path = ""
        self.init_ui()
    def closeEvent(self, event):
        clean_temp_files()
        event.accept()
    def init_ui(self):
        self.setWindowTitle("PDF/Изображение в DOCX")
        self.setMinimumSize(600, 400)
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        boxes_layout = QHBoxLayout()
        self.input_box = QGroupBox("Выберите файл")
        input_layout = QVBoxLayout()
        self.input_box.setLayout(input_layout)
        self.drop_area = DropArea(self)
        self.drop_area.setMinimumSize(200, 200)
        input_layout.addWidget(self.drop_area)
        self.select_file_btn = QPushButton("Выбрать файл")
        self.select_file_btn.setIcon(self.style().standardIcon(self.style().StandardPixmap.SP_DialogOpenButton))
        self.select_file_btn.clicked.connect(self.select_files)
        input_layout.addWidget(self.select_file_btn)
        self.output_box = QGroupBox("Выберите путь сохранения")
        output_layout = QVBoxLayout()
        self.output_box.setLayout(output_layout)
        icon_label = QLabel()
        icon = self.style().standardIcon(self.style().StandardPixmap.SP_FileDialogNewFolder)
        icon_label.setPixmap(icon.pixmap(64, 64))
        icon_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        output_layout.addWidget(icon_label)
        self.output_path_label = QLabel("Путь не выбран")
        self.output_path_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.output_path_label.setWordWrap(True)
        output_layout.addWidget(self.output_path_label)
        self.select_output_btn = QPushButton("Выбрать путь")
        self.select_output_btn.setIcon(self.style().standardIcon(self.style().StandardPixmap.SP_DialogSaveButton))
        self.select_output_btn.clicked.connect(self.select_output_path)
        output_layout.addWidget(self.select_output_btn)
        boxes_layout.addWidget(self.input_box)
        boxes_layout.addWidget(self.output_box)
        main_layout.addLayout(boxes_layout)
        self.ocr_engine_combo = QComboBox()
        if PADDLE_OCR_AVAILABLE:
            self.ocr_engine_combo.addItem("PaddleOCR")
        if TROCR_AVAILABLE:
            self.ocr_engine_combo.addItem("TrOCR")
        self.ocr_engine_combo.hide()  
        self.convert_btn = QPushButton("Конвертировать")
        self.convert_btn.setEnabled(False)
        self.convert_btn.clicked.connect(self.process_files)
        self.convert_btn.setMinimumHeight(50)  
        self.convert_btn.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                font-size: 16px;
                font-weight: bold;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #0D47A1;
            }
            QPushButton:disabled {
                background-color: #BDBDBD;
            }
        """)
        main_layout.addWidget(self.convert_btn)
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        main_layout.addWidget(self.progress_bar)
        self.progress_bar.hide()  
        self.status_label = QLabel()
        main_layout.addWidget(self.status_label)
        self.status_label.hide()  
    def select_files(self):
        file_dialog = QFileDialog()
        file_paths, _ = file_dialog.getOpenFileNames(
            self,
            "Выберите файлы",
            "",
            "Все поддерживаемые форматы (*.pdf *.jpg *.jpeg *.png *.bmp);;PDF файлы (*.pdf);;Изображения (*.jpg *.jpeg *.png *.bmp)"
        )
        if file_paths:
            self.file_paths = file_paths
            self.drop_area.label.setText(f"Выбрано файлов: {len(file_paths)}")
            self.update_convert_button()
    def select_output_path(self):
        output_path, _ = QFileDialog.getSaveFileName(
            self,
            "Сохранить DOCX файл",
            "",
            "DOCX файлы (*.docx)"
        )
        if output_path:
            if not output_path.endswith('.docx'):
                output_path += '.docx'
            self.output_path = output_path
            path_parts = os.path.split(output_path)
            display_path = f".../{path_parts[1]}"
            self.output_path_label.setText(display_path)
            self.update_convert_button()
    def update_convert_button(self):
        self.convert_btn.setEnabled(bool(self.file_paths) and bool(self.output_path))
    def process_files(self):
        if not self.file_paths or not self.output_path:
            QMessageBox.warning(self, "Предупреждение", "Выберите файлы и путь сохранения")
            return
        self.progress_bar.show()
        self.progress_bar.setValue(0)
        self.select_file_btn.setEnabled(False)
        self.select_output_btn.setEnabled(False)
        self.convert_btn.setEnabled(False)
        ocr_engine = self.ocr_engine_combo.currentText()
        self.processing_thread = ProcessingThread(
            self.ocr_processor, 
            self.file_paths, 
            self.output_path,
            ocr_engine
        )
        self.processing_thread.progress_signal.connect(self.update_progress)
        self.processing_thread.finished_signal.connect(self.processing_finished)
        self.processing_thread.error_signal.connect(self.processing_error)
        self.processing_thread.start()
    def update_progress(self, value):
        self.progress_bar.setValue(value)
    def processing_finished(self, docx_path):
        self.progress_bar.setValue(100)
        self.select_file_btn.setEnabled(True)
        self.select_output_btn.setEnabled(True)
        self.convert_btn.setEnabled(True)
        reply = QMessageBox.question(
            self, 
            "Обработка завершена", 
            f"Файл успешно сохранен как:\n{self.output_path}\n\nОткрыть файл?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            os.startfile(self.output_path) if os.name == 'nt' else os.system(f'xdg-open "{self.output_path}"')
        self.progress_bar.hide()
    def processing_error(self, error_message):
        self.progress_bar.setValue(0)
        self.progress_bar.hide()
        self.select_file_btn.setEnabled(True)
        self.select_output_btn.setEnabled(True)
        self.convert_btn.setEnabled(True)
        QMessageBox.critical(self, "Ошибка", error_message)
def main():
    os.makedirs(os.path.join(os.path.dirname(os.path.abspath(__file__)), "models"), exist_ok=True)
    os.makedirs(os.path.join(os.path.dirname(os.path.abspath(__file__)), "temp"), exist_ok=True)
    app = QApplication(sys.argv)
    app.setApplicationName("PDF/Изображение в DOCX")
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
if __name__ == "__main__":
    main() 
